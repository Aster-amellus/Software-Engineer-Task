"""Task orchestration for the arXiv review pipeline.

This module follows the staged flow described in the SRS: keyword expansion,
arXiv search, dedup, TopK selection, download, parse, chunk, embed, retrieve,
structured extraction, write, and export. The implementation keeps
dependencies lightweight (mock-friendly) while persisting progress to the
database and streaming events via Redis Pub/Sub for the WebSocket layer.
"""

from __future__ import annotations

import time
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import List

from sqlalchemy.orm import Session

try:  # optional heavy deps; fall back to lightweight behavior in tests
    import chromadb
    from chromadb.config import Settings as ChromaSettings
except ImportError:  # pragma: no cover - triggered in minimal CI envs
    chromadb = None
    ChromaSettings = None

try:
    import fitz
except ImportError:  # pragma: no cover
    fitz = None

from core.config import get_settings
from infrastructure.arxiv import ArxivAdapter, MockArxivAdapter, PaperMetadata
from infrastructure.embedding import EmbeddingProvider, MockEmbeddingProvider
from infrastructure.llm import LLMProvider, MockLLMProvider, resolve_llm_provider
from infrastructure.pubsub import publish_project_event
from models import Analysis, Export, Paper, Project
from utils.files import project_storage_path

STAGES = [
    "KEYWORD_EXPAND",
    "ARXIV_SEARCH",
    "DEDUP",
    "TOPK_SELECT",
    "DOWNLOAD",
    "PARSE",
    "CHUNK",
    "EMBED",
    "RETRIEVE",
    "EXTRACT",
    "WRITE",
    "EXPORT",
    "DONE",
]


@dataclass
class PipelineConfig:
    search: dict
    runtime: dict
    providers: dict


class PipelineService:
    def __init__(
        self,
        db: Session,
        arxiv_adapter: ArxivAdapter | None = None,
        llm_provider: LLMProvider | None = None,
        embed_provider: EmbeddingProvider | None = None,
    ):
        self.db = db
        self.settings = get_settings()
        self.arxiv = arxiv_adapter or MockArxivAdapter()
        self.llm = llm_provider or MockLLMProvider()
        self.embed = embed_provider or MockEmbeddingProvider()

    # ---- helpers -----------------------------------------------------
    def _update_stage(self, project: Project, stage: str, progress: int, log: str | None = None):
        project.stage = stage
        project.progress = progress
        project.updated_at = datetime.utcnow()
        self.db.add(project)
        self.db.commit()
        payload = {"stage": stage, "progress": progress}
        if log:
            payload["log"] = log
        publish_project_event(project.id, "stage", payload)
        if log:
            publish_project_event(project.id, "log", {"message": log})

    def _log(self, project: Project, message: str):
        publish_project_event(project.id, "log", {"message": message})

    def _load_config(self, project: Project) -> PipelineConfig:
        cfg = project.config or {}
        return PipelineConfig(
            search=cfg.get("search", {}),
            runtime=cfg.get("runtime", {}),
            providers=cfg.get("providers", {}),
        )

    # ---- stages ------------------------------------------------------
    def run(self, project: Project):
        project.status = "running"
        self.db.commit()

        config = self._load_config(project)
        self.llm = resolve_llm_provider(config.providers.get("llm", {}).get("name"))

        try:
            keywords = self._keyword_expand(project)
            search_results = self._arxiv_search(project, keywords, config)
            deduped = self._dedup(project, search_results)
            topk = self._topk(project, deduped, config)
            downloads = self._download(project, topk, config)
            parsed = self._parse(project, downloads)
            chunks = self._chunk(project, parsed)
            self._embed(project, chunks)
            self._retrieve(project, chunks)
            self._extract(project, topk)
            self._write(project)
            self._export(project)
            project.status = "completed"
            project.progress = 100
            project.stage = "DONE"
            self.db.commit()
            publish_project_event(project.id, "done", {"status": project.status})
        except Exception as exc:  # pragma: no cover - defensive catch
            project.status = "failed"
            project.stage = "DONE"
            project.progress = project.progress or 0
            self.db.commit()
            publish_project_event(project.id, "error", {"message": str(exc)})
            raise

    def _keyword_expand(self, project: Project) -> list[str]:
        self._update_stage(project, "KEYWORD_EXPAND", 5, log="Expanding keywords")
        kw = project.keywords or []
        expanded = kw + [f"{project.topic} review"]
        project.keywords = expanded
        self.db.commit()
        return expanded

    def _arxiv_search(self, project: Project, keywords: list[str], cfg: PipelineConfig) -> List[PaperMetadata]:
        self._update_stage(project, "ARXIV_SEARCH", 15, log="Searching arXiv")
        query = " + ".join(keywords) if keywords else project.topic
        max_results = int(cfg.search.get("max_results", 10))
        start = int(cfg.search.get("start", 0))
        sort_by = cfg.search.get("sortBy", "submittedDate")
        sort_order = cfg.search.get("sortOrder", "descending")
        results = self.arxiv.search(query, start, max_results, sort_by=sort_by, sort_order=sort_order)
        self._log(project, f"Retrieved {len(results)} papers")
        return results

    def _dedup(self, project: Project, results: List[PaperMetadata]) -> List[PaperMetadata]:
        self._update_stage(project, "DEDUP", 25, log="Deduplicating papers")
        seen = set()
        deduped = []
        for meta in results:
            if meta.arxiv_id in seen:
                continue
            seen.add(meta.arxiv_id)
            deduped.append(meta)
        self._materialize_papers(project, deduped)
        return deduped

    def _topk(self, project: Project, papers: List[PaperMetadata], cfg: PipelineConfig) -> List[Paper]:
        self._update_stage(project, "TOPK_SELECT", 35, log="Selecting TopK papers")
        k = int(cfg.runtime.get("top_k", 6))
        selected = papers[:k]
        db_papers = (
            self.db.query(Paper)
            .filter(Paper.project_id == project.id, Paper.arxiv_id.in_([p.arxiv_id for p in selected]))
            .all()
        )
        return db_papers

    def _download(self, project: Project, papers: List[Paper], cfg: PipelineConfig) -> List[Paper]:
        self._update_stage(project, "DOWNLOAD", 45, log="Downloading PDFs")
        runtime_retry = int(cfg.runtime.get("retry", 3))
        storage_dir = Path(project_storage_path(project.id))
        for paper in papers:
            pdf_path = storage_dir / f"{paper.arxiv_id}.pdf"
            for attempt in range(runtime_retry):
                try:
                    self._write_mock_pdf(pdf_path, paper)
                    paper.local_path = str(pdf_path)
                    paper.download_status = "ok"
                    break
                except Exception:  # pragma: no cover - keep retry path minimal
                    time.sleep(0.1)
                    paper.download_status = "failed"
            self.db.add(paper)
        self.db.commit()
        return papers

    def _parse(self, project: Project, papers: List[Paper]) -> dict[int, str]:
        self._update_stage(project, "PARSE", 55, log="Parsing PDFs")
        parsed: dict[int, str] = {}
        for paper in papers:
            if not paper.local_path:
                continue
            if fitz:
                with fitz.open(paper.local_path) as doc:
                    text = "\n".join(page.get_text() for page in doc)
            else:
                text = Path(paper.local_path).read_text(encoding="utf-8")
            parsed[paper.id] = text
        return parsed

    def _chunk(self, project: Project, parsed: dict[int, str]) -> dict[int, list[str]]:
        self._update_stage(project, "CHUNK", 65, log="Chunking text")
        chunks: dict[int, list[str]] = {}
        for paper_id, text in parsed.items():
            tokens = list(text)
            size = 500
            overlap = 50
            slices = []
            for start in range(0, len(tokens), size - overlap):
                slices.append("".join(tokens[start : start + size]))
            chunks[paper_id] = slices or [text]
        return chunks

    def _embed(self, project: Project, chunks: dict[int, list[str]]):
        self._update_stage(project, "EMBED", 75, log="Embedding and persisting to Chroma")
        if not chromadb:
            self._log(project, "Chroma not installed; skipping persistence")
            return

        client = chromadb.PersistentClient(
            path=self.settings.chroma_dir,
            settings=ChromaSettings(anonymized_telemetry=False),
        )
        collection = client.get_or_create_collection(f"project_{project.id}")
        ids: list[str] = []
        texts: list[str] = []
        metadatas: list[dict] = []
        for paper_id, chunk_list in chunks.items():
            for idx, chunk in enumerate(chunk_list):
                ids.append(f"paper:{paper_id}:chunk:{idx}")
                texts.append(chunk)
                metadatas.append({"paper_id": paper_id, "chunk": idx})
        if texts:
            vectors = self.embed.embed(texts)
            collection.add(ids=ids, documents=texts, metadatas=metadatas, embeddings=vectors)

    def _retrieve(self, project: Project, chunks: dict[int, list[str]]):
        self._update_stage(project, "RETRIEVE", 80, log="Retrieving evidence")
        # For MVP, retrieval is a no-op; embeddings are persisted and query can run later.

    def _extract(self, project: Project, papers: List[Paper]):
        self._update_stage(project, "EXTRACT", 88, log="Running structured extraction")
        for paper in papers:
            prompt = paper.abstract or paper.title or ""
            extracted = self.llm.generate_structured(prompt, schema={})
            analysis = Analysis(
                project_id=project.id,
                paper_id=paper.id,
                schema_version="v1",
                extracted=extracted,
                summary=extracted.get("limitations", ""),
                token_cost=10,
            )
            self.db.add(analysis)
        self.db.commit()

    def _write(self, project: Project):
        self._update_stage(project, "WRITE", 92, log="Writing markdown report")
        evidence = [paper.abstract or paper.title for paper in project.papers]
        outline = f"Literature review for {project.topic}"
        project.report_markdown = self.llm.write_markdown(outline, [e for e in evidence if e])
        self.db.commit()

    def _export(self, project: Project):
        self._update_stage(project, "EXPORT", 97, log="Exporting report")
        storage_dir = Path(project_storage_path(project.id))
        storage_dir.mkdir(parents=True, exist_ok=True)

        md_path = storage_dir / "report.md"
        md_path.write_text(project.report_markdown or "", encoding="utf-8")
        self._persist_export(project, "md", md_path)

        # DOCX export
        try:
            from docx import Document

            doc = Document()
            doc.add_heading(project.topic, 0)
            for paragraph in (project.report_markdown or "").split("\n\n"):
                if paragraph.strip():
                    doc.add_paragraph(paragraph)
            docx_path = storage_dir / "report.docx"
            doc.save(docx_path)
            self._persist_export(project, "docx", docx_path)
        except Exception:  # pragma: no cover
            self._log(project, "DOCX export failed; continuing")

        # PDF export via WeasyPrint
        try:
            from weasyprint import HTML

            html_content = f"<h1>{project.topic}</h1>" + "".join(
                f"<p>{p}</p>" for p in (project.report_markdown or "").split("\n")
            )
            pdf_path = storage_dir / "report.pdf"
            HTML(string=html_content).write_pdf(pdf_path)
            self._persist_export(project, "pdf", pdf_path)
        except Exception:  # pragma: no cover
            self._log(project, "PDF export failed; continuing")

    def _persist_export(self, project: Project, fmt: str, path: Path):
        export = Export(project_id=project.id, format=fmt, local_path=str(path), status="ok")
        self.db.add(export)
        self.db.commit()

    # ---- utils -------------------------------------------------------
    def _materialize_papers(self, project: Project, results: List[PaperMetadata]):
        for meta in results:
            existing = (
                self.db.query(Paper)
                .filter(Paper.project_id == project.id, Paper.arxiv_id == meta.arxiv_id)
                .first()
            )
            if existing:
                continue
            paper = Paper(
                project_id=project.id,
                arxiv_id=meta.arxiv_id,
                title=meta.title,
                authors=meta.authors,
                abstract=meta.abstract,
                categories=meta.categories,
                published_at=meta.published_at,
                pdf_url=meta.pdf_url,
                download_status="pending",
            )
            self.db.add(paper)
        self.db.commit()
        self.db.refresh(project)

    @staticmethod
    def _write_mock_pdf(path: Path, paper: Paper):
        path.parent.mkdir(parents=True, exist_ok=True)
        if fitz:
            doc = fitz.open()
            page = doc.new_page()
            page.insert_text((72, 72), f"{paper.title}\n\n{paper.abstract or 'No abstract'}")
            doc.save(path)
            doc.close()
        else:
            path.write_text(f"{paper.title}\n\n{paper.abstract or 'No abstract'}", encoding="utf-8")
